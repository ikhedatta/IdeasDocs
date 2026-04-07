"""
DOCX Parser using python-docx.

RAGFlow Source: deepdoc/parser/docx_parser.py
Extracts paragraphs, headers, tables, and lists with structure preservation.
"""

from io import BytesIO
from docx import Document
from docx.table import Table
from chunkers.models import ContentBlock, BlockType
from .base import ParserRegistry


@ParserRegistry.register([".docx", ".doc"])
class DocxParser:
    """
    DOCX parser that preserves document structure.
    
    Extracts:
    - Paragraphs with style-based type detection (Heading 1-6 → HEADER)
    - Tables as structured text (preserving rows/columns)
    - Lists (bullet/numbered) as LIST blocks
    """

    def parse(
        self, file_bytes: bytes, filename: str, config: dict | None = None
    ) -> list[ContentBlock]:
        doc = Document(BytesIO(file_bytes))
        blocks: list[ContentBlock] = []
        element_order = 0

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # Paragraph
                para = self._find_paragraph(doc, element)
                if para and para.text.strip():
                    block_type = self._detect_style_type(para)
                    blocks.append(ContentBlock(
                        text=para.text.strip(),
                        block_type=block_type,
                        metadata={"style": para.style.name if para.style else "Normal"},
                    ))
                    element_order += 1

            elif tag == "tbl":
                # Table
                table = self._find_table(doc, element)
                if table:
                    table_text = self._table_to_text(table)
                    if table_text.strip():
                        blocks.append(ContentBlock(
                            text=table_text,
                            block_type=BlockType.TABLE,
                            metadata={"rows": len(table.rows), "cols": len(table.columns)},
                        ))
                        element_order += 1

        return blocks

    def _find_paragraph(self, doc: Document, element):
        """Find the Paragraph object corresponding to an XML element."""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table(self, doc: Document, element):
        """Find the Table object corresponding to an XML element."""
        for table in doc.tables:
            if table._element is element:
                return table
        return None

    def _detect_style_type(self, para) -> BlockType:
        """Detect block type from paragraph style."""
        style_name = para.style.name.lower() if para.style else ""
        
        if "heading" in style_name or "title" in style_name:
            return BlockType.HEADER
        if "list" in style_name or "bullet" in style_name:
            return BlockType.LIST
        if "code" in style_name:
            return BlockType.CODE
        
        # Check for bullet/number formatting
        text = para.text.strip()
        if text and (text[0] in "•●○◦-–" or 
                     (len(text) > 2 and text[0].isdigit() and text[1] in ".)")):
            return BlockType.LIST
        
        return BlockType.TEXT

    def _table_to_text(self, table: Table) -> str:
        """
        Convert table to markdown-style text representation.
        
        RAGFlow stores tables as structured text so they can be
        chunked and searched as a unit with surrounding context.
        """
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if len(rows) > 1:
            # Add markdown header separator after first row
            header_sep = "| " + " | ".join(["---"] * len(table.columns)) + " |"
            rows.insert(1, header_sep)

        return "\n".join(rows)
